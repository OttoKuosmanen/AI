import json
from docx import Document


relative_location_name = '../data/AI/GPT35_shorter.json'

def read(relative_location_name):
    with open(relative_location_name, 'r') as f:
        out = json.load(f)
        return out

def extract(analysis_data):
    out = []
    for i in analysis_data:
        out.append(i)
    return out


# reading in data
analysis_data = read(relative_location_name)

# Extracting answers
answers = extract(analysis_data)

document_answers = Document()

# Writing out questions and asnwers and saving it to the wordfile.


for index, answer in enumerate(answers):
    number = str(index + 1)
    document_answers.add_paragraph(number)
    document_answers.add_paragraph(answer)

document_answers.save("answers_GPT3_5.docx")



